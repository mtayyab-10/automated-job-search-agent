import os
import tempfile
import fitz  # PyMuPDF
import docx
from pathlib import Path

def extract_text(uploaded_file) -> str:
    """
    Accepts a Streamlit uploaded file object, saves it temporarily to disk,
    detects whether it is a PDF or DOCX, extracts all readable text using PyMuPDF (fitz)
    or python-docx, cleans up the temp file, and returns the extracted text.
    Raises ValueError if the file format is unsupported.
    """
    file_name = uploaded_file.name
    suffix = Path(file_name).suffix.lower()

    if suffix not in ('.pdf', '.docx'):
        raise ValueError(f"Unsupported file type: {suffix}. Only PDF and DOCX files are supported.")

    # Save uploaded file to a temporary file
    temp_fd, temp_path = tempfile.mkstemp(suffix=suffix)
    try:
        with os.fdopen(temp_fd, 'wb') as temp_file:
            temp_file.write(uploaded_file.getvalue())

        extracted_text = ""
        if suffix == '.pdf':
            doc = fitz.open(temp_path)
            pages_text = []
            for page in doc:
                pages_text.append(page.get_text())
            extracted_text = "\n".join(pages_text)
            doc.close()
        elif suffix == '.docx':
            doc = docx.Document(temp_path)
            paragraphs_text = [para.text for para in doc.paragraphs]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        paragraphs_text.append(cell.text)
            extracted_text = "\n".join(paragraphs_text)

        print(f"Extracted {len(extracted_text)} characters from {file_name}")
        return extracted_text

    finally:
        # Clean up temp file
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception as e:
                print(f"Error removing temporary file {temp_path}: {e}")
